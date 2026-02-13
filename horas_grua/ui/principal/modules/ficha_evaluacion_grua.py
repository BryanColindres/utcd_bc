import customtkinter as ctk
from tkinter import Toplevel, messagebox
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from datetime import datetime
import os, sys

sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "../../..")))
from db_controller import actualizar_completado_orden

# Recomiendo inicializar el theme en la app principal:
# ctk.set_appearance_mode("Light")
# ctk.set_default_color_theme("blue")


class EvaluacionProveedorModal(Toplevel):
    """
    Modal para evaluación de proveedor de grúa.
    Ajustes: ventana más grande y textos de opciones en negro y con mayor tamaño.
    """

    def __init__(self, parent, orden, proveedor, supervisor, horas, equipo, callback=None):
        super().__init__(parent)
        self.title("Evaluación del Proveedor de Grúa")
        # Tamaño aumentado
        self.geometry("1000x980")
        self.transient(parent)
        self.grab_set()
        self.configure(bg="#FFFFFF")
        self.supervisor = supervisor
        self.callback = callback

        # Estilos
        self._label_font = ("TT NORMS PRO", 15)
        self._title_font = ("TT NORMS PRO", 22, "bold")
        self._input_width = 420
        self._entry_fg = "#FFFFFF"
        self._entry_text_color = "#191616"
        self._option_font = ("TT NORMS PRO", 14)     # fuente para opciones (radio)
        self._option_text_color = "#191616"         # color negro para texto de opciones
        self._accent = "#002D86"

        # Título
        ctk.CTkLabel(self, text="FICHA DE EVALUACIÓN DE SERVICIO DE GRÚA",
                     font=self._title_font, text_color="#070707").pack(pady=(18, 8))

        form = ctk.CTkScrollableFrame(
            self,
            fg_color="white",
            corner_radius=10,
            width=960,
            height=820
        )
        form.pack(fill="both", expand=True, padx=20, pady=12)

        form.grid_columnconfigure(0, weight=0, minsize=360)
        form.grid_columnconfigure(1, weight=1)

        def add_label(text, row, sticky="w"):
            ctk.CTkLabel(form, text=text, font=self._label_font, text_color="#2B2B2B").grid(
                row=row, column=0, sticky=sticky, padx=(12, 6), pady=10
            )

        # 1. Servicio tercerizado
        add_label("1. ¿Cuenta su sector con servicio de grúa tercerizado?", 0)
        self.servicio_tercerizado = ctk.StringVar(value="NO")
        rb_frame = ctk.CTkFrame(form, fg_color="transparent")
        rb_frame.grid(row=0, column=1, sticky="w", padx=6)
        ctk.CTkRadioButton(rb_frame, text="SI", variable=self.servicio_tercerizado, value="SI",
                           font=self._option_font, text_color=self._option_text_color).pack(side="left", padx=(0,10))
        ctk.CTkRadioButton(rb_frame, text="NO", variable=self.servicio_tercerizado, value="NO",
                           font=self._option_font, text_color=self._option_text_color).pack(side="left")

        # 2. Nombre del proveedor (autollenar y bloqueado)
        add_label("2. Nombre del proveedor:", 1)
        self.nombre_proveedor = ctk.CTkEntry(form, width=self._input_width,
                                             fg_color=self._entry_fg, text_color=self._entry_text_color,
                                             corner_radius=6)
        self.nombre_proveedor.grid(row=1, column=1, sticky="w", padx=6, pady=6)
        self.nombre_proveedor.insert(0, proveedor or "")
        self.nombre_proveedor.configure(state="disabled")

        # 3. Orden y horas aprobadas (autollenar y bloqueado)
        add_label("3. Orden de compra y horas aprobadas:", 2)
        self.orden_info = ctk.CTkEntry(form, width=self._input_width,
                                       fg_color=self._entry_fg, text_color=self._entry_text_color,
                                       corner_radius=6)
        self.orden_info.grid(row=2, column=1, sticky="w", padx=6, pady=6)
        self.orden_info.insert(0, f"{orden} - Horas aprobadas: {horas}")
        self.orden_info.configure(state="disabled")
        
        self.orden_compra = orden
        self.tipo_equipo_asignado = equipo
        self.nombre_proveedor_val = proveedor
        # 3.1 Tipo equipo (autollenar y bloqueado)
        add_label("3.1 Tipo de equipo asignado:", 3)
        self.tipo_equipo = ctk.CTkEntry(form, width=self._input_width,
                                        fg_color=self._entry_fg, text_color=self._entry_text_color,
                                        corner_radius=6)
        self.tipo_equipo.grid(row=3, column=1, sticky="w", padx=6, pady=6)
        self.tipo_equipo.insert(0, equipo or "")
        self.tipo_equipo.configure(state="disabled")

        # 4. Disponibilidad inmediata
        add_label("4. ¿El proveedor brinda disponibilidad inmediata?", 4)
        self.disponible = ctk.StringVar(value="NO")
        rb_frame2 = ctk.CTkFrame(form, fg_color="transparent")
        rb_frame2.grid(row=4, column=1, sticky="w", padx=6)
        ctk.CTkRadioButton(rb_frame2, text="SI", variable=self.disponible, value="SI",
                           font=self._option_font, text_color=self._option_text_color).pack(side="left", padx=(0,10))
        ctk.CTkRadioButton(rb_frame2, text="NO", variable=self.disponible, value="NO",
                           font=self._option_font, text_color=self._option_text_color).pack(side="left")

        # 5. Grúa sin fallas
        add_label("5. ¿La grúa ha presentado fallas durante el servicio?", 5)
        self.sin_fallas = ctk.StringVar(value="NO")
        rb_frame3 = ctk.CTkFrame(form, fg_color="transparent")
        rb_frame3.grid(row=5, column=1, sticky="w", padx=6)
        ctk.CTkRadioButton(rb_frame3, text="SI", variable=self.sin_fallas, value="SI",
                           font=self._option_font, text_color=self._option_text_color).pack(side="left", padx=(0,10))
        ctk.CTkRadioButton(rb_frame3, text="NO", variable=self.sin_fallas, value="NO",
                           font=self._option_font, text_color=self._option_text_color).pack(side="left")

        # 6. Veces que NO brindó el servicio
        add_label("6. Veces que NO brindó el servicio:", 6)
        self.veces_no = ctk.CTkEntry(form, width=140,
                                     fg_color=self._entry_fg, text_color=self._entry_text_color,
                                     corner_radius=6)
        self.veces_no.grid(row=6, column=1, sticky="w", padx=6, pady=6)

        # 7. Cumple especificaciones técnicas
        add_label("7. ¿El equipo cumple especificaciones técnicas?", 7)
        self.especificaciones = ctk.StringVar(value="NO")
        rb_frame4 = ctk.CTkFrame(form, fg_color="transparent")
        rb_frame4.grid(row=7, column=1, sticky="w", padx=6)
        ctk.CTkRadioButton(rb_frame4, text="SI", variable=self.especificaciones, value="SI",
                           font=self._option_font, text_color=self._option_text_color).pack(side="left", padx=(0,10))
        ctk.CTkRadioButton(rb_frame4, text="NO", variable=self.especificaciones, value="NO",
                           font=self._option_font, text_color=self._option_text_color).pack(side="left")

        # 8. Calificación
        add_label("8. Califique el servicio:", 8)
        rb_frame5 = ctk.CTkFrame(form, fg_color="transparent")
        rb_frame5.grid(row=8, column=1, sticky="w", padx=6)
        self.calificacion = ctk.StringVar(value="Regular")
        ctk.CTkRadioButton(rb_frame5, text="Muy Bueno", variable=self.calificacion, value="Muy Bueno",
                           font=self._option_font, text_color=self._option_text_color).pack(side="left", padx=(0,10))
        ctk.CTkRadioButton(rb_frame5, text="Regular", variable=self.calificacion, value="Regular",
                           font=self._option_font, text_color=self._option_text_color).pack(side="left", padx=(0,10))
        ctk.CTkRadioButton(rb_frame5, text="Mal Servicio", variable=self.calificacion, value="Mal Servicio",
                           font=self._option_font, text_color=self._option_text_color).pack(side="left")

        # 9. Observaciones (textbox) con fondo claro
        add_label("9. Observaciones:", 9, sticky="nw")
        self.observaciones = ctk.CTkTextbox(form, width=720, height=120, fg_color="#FBFBFB", text_color="#191616")
        self.observaciones.grid(row=9, column=1, sticky="w", padx=6, pady=8)

        # Botones inferiores (texto en negro)
        btn_frame = ctk.CTkFrame(form, fg_color="transparent")
        btn_frame.grid(row=10, column=0, columnspan=2, pady=(14, 18))
        ctk.CTkButton(btn_frame, text="Guardar Evaluación", fg_color="#5FD0DF", hover_color="#56C0CE",
                      command=self.guardar, width=220, height=44, text_color="#000000",
                      font=("TT NORMS PRO", 13, "bold")).pack(side="left", padx=10)
        ctk.CTkButton(btn_frame, text="Cancelar", fg_color="#F1F1F4", hover_color="#EDEDED",
                      command=self._cancel, width=140, height=44, text_color="#000000",
                      font=("TT NORMS PRO", 13)).pack(side="left", padx=10)

    # -------------------------------------------------------------
    def _cancel(self):
        if messagebox.askyesno("Cancelar", "¿Desea cerrar sin guardar?"):
            self.destroy()

    # -------------------------------------------------------------
    def generar_word(self, datos):
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
        style.font.size = Pt(11)

        titulo = doc.add_heading("FICHA DE EVALUACIÓN DE PROVEEDOR DE SERVICIO DE GRÚA", level=1)
        titulo.alignment = 1
        doc.add_paragraph("\nSe requiere marcar con un '☑' la opción a escoger\n")

        def opcion(valor, esperado):
            return "☑" if valor == esperado else "☐"

        def texto_subrayado(parrafo, texto):
            run = parrafo.add_run(texto)
            run.font.underline = True
            run.font.color.rgb = RGBColor(0, 0, 180)
            run.font.size = Pt(11)

        p = doc.add_paragraph()
        p.add_run("1. ¿Cuenta su sector con servicio de grúa tercerizado?   ")
        p.add_run(f"SI {opcion(datos['servicio_tercerizado'], 'SI')}   NO {opcion(datos['servicio_tercerizado'], 'NO')}")

        p = doc.add_paragraph("2. Nombre del proveedor:  ")
        texto_subrayado(p, datos.get('nombre_proveedor', '') or " ")

        p = doc.add_paragraph("3. Orden de compra actual y horas aprobadas:  ")
        texto_subrayado(p, datos.get('orden_info', '') or " ")

        p = doc.add_paragraph("3.1 Tipo de equipo asignado:  ")
        texto_subrayado(p, datos.get('tipo_equipo', '') or " ")

        p = doc.add_paragraph()
        p.add_run("4. ¿El proveedor brinda disponibilidad inmediata?   ")
        p.add_run(f"SI {opcion(datos.get('disponible'), 'SI')}   NO {opcion(datos.get('disponible'), 'NO')}")

        p = doc.add_paragraph()
        p.add_run("5. ¿Proveedor no ha prestado el servicio por grúa en mal estado?   ")
        p.add_run(f"SI {opcion(datos.get('sin_fallas'), 'SI')}   NO {opcion(datos.get('sin_fallas'), 'NO')}")

        p = doc.add_paragraph("6. Veces que NO brindó el servicio:  ")
        texto_subrayado(p, datos.get('veces_no', '') or " ")

        p = doc.add_paragraph()
        p.add_run("7. ¿El equipo cumple especificaciones técnicas?   ")
        p.add_run(f"SI {opcion(datos.get('especificaciones'), 'SI')}   NO {opcion(datos.get('especificaciones'), 'NO')}")

        p = doc.add_paragraph()
        p.add_run("8. Calificación del servicio:   ")
        p.add_run(f"MUY BUENO {opcion(datos.get('calificacion'), 'Muy Bueno')}   ")
        p.add_run(f"REGULAR {opcion(datos.get('calificacion'), 'Regular')}   ")
        p.add_run(f"MAL SERVICIO {opcion(datos.get('calificacion'), 'Mal Servicio')}")

        doc.add_paragraph("\n9. Observaciones:")
        p = doc.add_paragraph()
        texto_subrayado(p, datos.get("observaciones", "") or " ")

        doc.add_paragraph(f"\n\nNombre: {self.supervisor[0]}")
        
        p = doc.add_paragraph("Fecha:  ")
        texto_subrayado(p, datetime.now().strftime("%d/%m/%Y"))

        nombre_archivo = f"Evaluacion_{self.nombre_proveedor_val}_{self.orden_compra}_{self.tipo_equipo_asignado}.docx"
        doc.save(nombre_archivo)
        return nombre_archivo

    # -------------------------------------------------------------
    def guardar(self):
        datos = {
            "servicio_tercerizado": self.servicio_tercerizado.get(),
            "nombre_proveedor": self.nombre_proveedor.get(),
            "orden_info": self.orden_info.get(),
            "disponible": self.disponible.get(),
            "sin_fallas": self.sin_fallas.get(),
            "veces_no": self.veces_no.get(),
            "especificaciones": self.especificaciones.get(),
            "calificacion": self.calificacion.get(),
            "observaciones": self.observaciones.get("0.0", "end").strip(),
            "tipo_equipo": self.tipo_equipo.get()
        }

        try:
            # Generar el documento
            archivo = self.generar_word(datos)

            # Intento actualizar BD (igual que antes)
            actualizado = False
            try:
                actualizado = actualizar_completado_orden(datos)
            except Exception as e_upd:
                # Si falla la actualización, lo notificamos pero NO rompemos el flujo.
                # Se devuelve actualizado = False para que el caller decida.
                print(f"[EvaluacionProveedorModal] Error actualizando orden: {e_upd}")
                actualizado = False

            # Notificar al usuario del archivo generado
            if archivo:
                messagebox.showinfo("Documento creado", f"El archivo '{archivo}' ha sido generado correctamente.")

            # Devolver resultado al caller mediante callback (si se pasó)
            if self.callback:
                # enviar un dict con los datos, nombre de archivo y flag de actualización
                try:
                    self.callback({
                        "datos": datos,
                        "archivo": archivo,
                        "actualizado": bool(actualizado)
                    })
                except Exception as e_cb:
                    print(f"[EvaluacionProveedorModal] Error en callback: {e_cb}")

            # Cerrar modal
            self.destroy()

        except Exception as e:
            messagebox.showerror("Error", f"No se pudo guardar: {str(e)}")
            # Intentamos devolver info al callback aunque haya fallado
            if self.callback:
                try:
                    self.callback({"datos": datos, "archivo": None, "actualizado": False})
                except:
                    pass
            return
